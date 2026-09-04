"""
Tests de la lecture universelle des fichiers .txt et .docx sans tableau.

Valide les 3 stratégies de lecture en cascade : délimité, colonnes
alignées par espaces, blocs clé-valeur — ainsi que le rejet propre d'un
texte réellement non structuré.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from ingestion.ingest import load_file, IngestionError


def test_txt_delimited(tmp_path):
    content = (
        "Username|System|Account Status|Employee Status\n"
        "jkonan|Active Directory|Active|Active\n"
        "bafolabi|SAP|Active|Terminated\n"
    )
    path = tmp_path / "delimited.txt"
    path.write_text(content, encoding="utf-8")

    df = load_file(path)
    assert len(df) == 2
    assert "username" in df.columns
    assert "system" in df.columns
    print("OK - test_txt_delimited")


def test_txt_fixed_width(tmp_path):
    content = (
        "Username    System              Account Status    Employee Status\n"
        "mkeita      Active Directory    Active            Active\n"
        "pyao        VPN                 Active            Terminated\n"
    )
    path = tmp_path / "fixedwidth.txt"
    path.write_text(content, encoding="utf-8")

    df = load_file(path)
    assert len(df) == 2
    assert "username" in df.columns
    print("OK - test_txt_fixed_width")


def test_txt_key_value_blocks(tmp_path):
    content = (
        "Nom d'utilisateur: sagbato\n"
        "Systeme: SIEM ArcSight\n"
        "Statut compte: Actif\n"
        "Statut RH: Actif\n"
        "\n"
        "Nom d'utilisateur: rzongo\n"
        "Systeme: CRM\n"
        "Statut compte: Actif\n"
        "Statut RH: Parti\n"
    )
    path = tmp_path / "keyvalue.txt"
    path.write_text(content, encoding="utf-8")

    df = load_file(path)
    assert len(df) == 2
    assert "username" in df.columns
    assert set(df["username"]) == {"sagbato", "rzongo"}
    print("OK - test_txt_key_value_blocks")


def test_txt_unreadable_raises_clear_error(tmp_path):
    content = (
        "Ceci est un simple paragraphe de texte libre, sans structure "
        "reconnaissable, comme le contenu d'un email ou d'une note."
    )
    path = tmp_path / "unreadable.txt"
    path.write_text(content, encoding="utf-8")

    try:
        load_file(path)
        assert False, "Une IngestionError aurait dû être levée"
    except IngestionError as e:
        assert "Impossible d'interpréter" in str(e)
        print("OK - test_txt_unreadable_raises_clear_error")


def test_docx_freetext_keyvalue(tmp_path):
    """.docx sans aucun tableau, données en paragraphes clé-valeur."""
    from docx import Document

    doc = Document()
    doc.add_heading("Fiches comptes", level=1)
    for username, system, status in [
        ("dyao", "Active Directory", "Parti"),
        ("kboni", "SAP", "Actif"),
    ]:
        doc.add_paragraph(f"Nom d'utilisateur: {username}")
        doc.add_paragraph(f"Systeme: {system}")
        doc.add_paragraph(f"Statut RH: {status}")
        doc.add_paragraph("")

    path = tmp_path / "freetext.docx"
    doc.save(str(path))

    df = load_file(path)
    assert len(df) == 2
    assert "username" in df.columns
    print("OK - test_docx_freetext_keyvalue")


if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        test_txt_delimited(tmp_path)
        test_txt_fixed_width(tmp_path)
        test_txt_key_value_blocks(tmp_path)
        test_txt_unreadable_raises_clear_error(tmp_path)
        test_docx_freetext_keyvalue(tmp_path)
    print("\nTous les tests sont passés.")


def test_missing_system_column_falls_back_to_filename():
    """
    Un export brut d'un seul système (ex. extraction AD pure) sans colonne
    'system' ne doit plus faire échouer l'ingestion : le nom du fichier
    sert de repli automatique.
    """
    import tempfile
    from ingestion.ingest import load_file

    content = "SAM Account Name,Display Name,Account Status\nadmin_test,Compte Test,Active\n"
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".csv", delete=False, prefix="ActiveDirectory_export_"
    ) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    df = load_file(tmp_path)
    assert "system" in df.columns
    assert df.loc[0, "system"] == Path(tmp_path).stem
    print(f"OK - test_missing_system_column_falls_back_to_filename (system={df.loc[0, 'system']!r})")


def test_missing_system_column_uses_explicit_default():
    """Un nom de système fourni explicitement prime sur le nom de fichier."""
    import tempfile
    from ingestion.ingest import load_file

    content = "SAM Account Name,Display Name,Account Status\nadmin_test,Compte Test,Active\n"
    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    df = load_file(tmp_path, default_system="Active Directory")
    assert df.loc[0, "system"] == "Active Directory"
    print("OK - test_missing_system_column_uses_explicit_default")


def test_duplicate_mapped_columns_are_merged_not_duplicated():
    """
    Deux colonnes sources distinctes qui pointent vers le même champ
    standard (ex. 'SAM Account Name' et 'Logon Name' -> toutes deux
    'username' dans un export AD) ne doivent jamais produire deux colonnes
    de même nom après standardisation — ça fait planter l'affichage en
    aval (Streamlit/Arrow). Elles doivent être fusionnées.
    """
    import pandas as pd
    from ingestion.ingest import standardize_columns

    df = pd.DataFrame({
        "SAM Account Name": ["jdupont", "kbrou"],
        "Logon Name": ["jdupont@corp.com", None],
        "Display Name": ["Jean Dupont", "Konan Brou"],
    })
    result = standardize_columns(df)
    assert len(result.columns) == len(set(result.columns)), "Colonnes dupliquées détectées"
    assert list(result["username"]) == ["jdupont", "kbrou"]
    print("OK - test_duplicate_mapped_columns_are_merged_not_duplicated")
