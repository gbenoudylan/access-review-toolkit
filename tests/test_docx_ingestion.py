"""
Tests du support des fichiers Word (.docx) en entrée.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from ingestion.ingest import load_file, IngestionError
from analysis.access_review import analyze_access


def _build_test_docx(path: Path, with_table: bool = True) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Export de test", level=1)
    doc.add_paragraph("Texte d'introduction, comme dans un vrai document reçu par email.")

    if with_table:
        headers = ["Identifiant", "Systeme", "Statut compte", "Statut RH"]
        rows = [
            ["user1", "Active Directory", "Actif", "Actif"],
            ["user2", "SAP", "Actif", "Parti"],
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row_data in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row_data):
                cells[i].text = val

    doc.add_paragraph("Note de bas de page.")
    doc.save(str(path))


def test_docx_with_table_is_ingested(tmp_path):
    docx_path = tmp_path / "test_with_table.docx"
    _build_test_docx(docx_path, with_table=True)

    df = load_file(docx_path)
    assert len(df) == 2
    assert "username" in df.columns
    assert "system" in df.columns
    print(f"OK - test_docx_with_table_is_ingested ({len(df)} lignes)")


def test_docx_without_table_raises_clear_error(tmp_path):
    docx_path = tmp_path / "test_no_table.docx"
    _build_test_docx(docx_path, with_table=False)

    try:
        load_file(docx_path)
        assert False, "Une IngestionError aurait dû être levée"
    except IngestionError as e:
        assert "tableau" in str(e).lower()
        print("OK - test_docx_without_table_raises_clear_error")


def test_docx_end_to_end_with_analysis(tmp_path):
    """Vérifie que le pipeline complet (ingestion + analyse) fonctionne sur un .docx."""
    docx_path = tmp_path / "test_full.docx"
    _build_test_docx(docx_path, with_table=True)

    df = load_file(docx_path)
    result = analyze_access(df)
    assert result.loc[1, "is_terminated_but_active"] == True  # user2 : Actif + Parti
    print("OK - test_docx_end_to_end_with_analysis")


if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        test_docx_with_table_is_ingested(tmp_path)
        test_docx_without_table_raises_clear_error(tmp_path)
        test_docx_end_to_end_with_analysis(tmp_path)
    print("\nTous les tests sont passés.")
