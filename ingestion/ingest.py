"""
Module d'ingestion universelle pour les exports d'accès/comptes.

Logique identique à celle validée sur le projet de gestion des
vulnérabilités : détection automatique de la ligne d'en-tête, gestion des
CSV mal formés, mapping des colonnes vers des noms standards internes.
Seul le référentiel de mapping (config/column_mapping.py) change de domaine.
"""

from __future__ import annotations
import logging
import re
from pathlib import Path

import pandas as pd

try:
    from rapidfuzz import fuzz
    _HAS_RAPIDFUZZ = True
except ImportError:
    _HAS_RAPIDFUZZ = False

from config.column_mapping import COLUMN_MAPPING, REQUIRED_FIELDS

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger("ingestion")


class IngestionError(Exception):
    """Erreur levée quand un fichier ne peut pas être exploité de façon fiable."""


def _detect_encoding(path: Path) -> str:
    """
    Détecte l'encodage texte d'un fichier plutôt que d'imposer l'UTF-8.

    Beaucoup d'exports réels (notamment depuis Excel ou des outils Windows)
    sont en Windows-1252/Latin-1, pas en UTF-8 — les lire en UTF-8 strict
    échoue ou corrompt les caractères accentués. On essaie plusieurs
    encodages courants dans l'ordre et on garde le premier qui décode le
    fichier sans erreur ; Latin-1 en dernier recours ne lève jamais
    d'erreur (il associe un caractère à chaque octet), donc la fonction
    retourne toujours un encodage utilisable.
    """
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            raw.decode(encoding)
            return encoding
        except UnicodeDecodeError:
            continue
    return "latin-1"  # filet de sécurité théorique, jamais atteint en pratique


def _normalize(text: str) -> str:
    return str(text).strip().lower().replace("_", " ").replace("-", " ")


def _score_header_row(row: pd.Series, column_mapping: dict = None) -> int:
    column_mapping = column_mapping or COLUMN_MAPPING
    all_variants = {
        _normalize(v) for variants in column_mapping.values() for v in variants
    }
    score = 0
    for cell in row:
        if pd.isna(cell):
            continue
        if _normalize(cell) in all_variants:
            score += 1
    return score


def _detect_header_row(raw: pd.DataFrame, column_mapping: dict = None, max_scan_rows: int = 15) -> int:
    best_row, best_score = 0, -1
    for i in range(min(max_scan_rows, len(raw))):
        score = _score_header_row(raw.iloc[i], column_mapping)
        if score > best_score:
            best_row, best_score = i, score
    if best_score <= 0:
        logger.warning("Aucune ligne d'en-tête reconnue, utilisation de la ligne 0.")
        return 0
    logger.info(f"En-tête détecté à la ligne {best_row} (score={best_score}).")
    return best_row


def _match_column(col_name: str, column_mapping: dict = None, threshold: int = 85) -> str | None:
    column_mapping = column_mapping or COLUMN_MAPPING
    col_norm = _normalize(col_name)

    for standard_name, variants in column_mapping.items():
        if col_norm in [_normalize(v) for v in variants]:
            return standard_name

    if _HAS_RAPIDFUZZ:
        best_field, best_score = None, 0
        for standard_name, variants in column_mapping.items():
            for v in variants:
                # token_sort_ratio : insensible à l'ordre des mots
                # (ex. "Statut compte" vs "Compte statut")
                s = fuzz.token_sort_ratio(col_norm, _normalize(v))
                if s > best_score:
                    best_field, best_score = standard_name, s
        if best_score >= threshold:
            logger.info(f"Colonne '{col_name}' -> '{best_field}' (fuzzy, score={best_score}).")
            return best_field

    return None


def standardize_columns(df: pd.DataFrame, column_mapping: dict = None) -> pd.DataFrame:
    """
    Renomme les colonnes reconnues vers leur nom standard.

    Il arrive qu'un même export contienne deux colonnes distinctes qui
    correspondent au même champ standard (ex. 'SAM Account Name' et
    'Logon Name' pointent toutes deux vers 'username' dans un export AD
    classique). Les mapper telles quelles produirait deux colonnes de même
    nom après renommage — invalide pour pandas/Arrow en aval (l'affichage
    Streamlit, notamment, lève une erreur sur des noms de colonnes
    dupliqués). On fusionne donc ces cas : la première colonne rencontrée
    fait foi, complétée par les valeurs non vides de la seconde là où elle
    a des trous, puis la seconde est supprimée.
    """
    rename_map, unmatched = {}, []
    claimed_by: dict[str, str] = {}  # nom standard -> colonne originale déjà utilisée

    for col in df.columns:
        matched = _match_column(col, column_mapping)
        if not matched:
            unmatched.append(col)
            continue
        if matched not in claimed_by:
            claimed_by[matched] = col
            rename_map[col] = matched
        else:
            primary_col = claimed_by[matched]
            df[primary_col] = df[primary_col].combine_first(df[col])
            df = df.drop(columns=[col])
            logger.info(
                f"Colonne '{col}' fusionnée dans '{primary_col}' (toutes deux -> '{matched}')."
            )

    if unmatched:
        logger.info(f"Colonnes non reconnues (ignorées) : {unmatched}")
    return df.rename(columns=rename_map)


def validate_required_fields(df: pd.DataFrame, required_fields: list = None) -> None:
    required_fields = required_fields if required_fields is not None else REQUIRED_FIELDS
    missing = [f for f in required_fields if f not in df.columns]
    if missing:
        raise IngestionError(
            f"Champs obligatoires manquants après mapping : {missing}. "
            f"Colonnes disponibles : {list(df.columns)}. "
            f"-> Ajoutez la variante manquante dans config/column_mapping.py"
        )


def _read_ragged_csv(path: Path) -> pd.DataFrame:
    import csv
    with open(path, newline="", encoding=_detect_encoding(path)) as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
        except csv.Error:
            dialect = csv.excel
        rows = list(csv.reader(f, dialect))
    max_cols = max(len(r) for r in rows) if rows else 0
    rows = [r + [None] * (max_cols - len(r)) for r in rows]
    return pd.DataFrame(rows)


def _try_delimited(lines: list[str], column_mapping: dict = None) -> pd.DataFrame | None:
    """
    Tentative n°1 : le texte est en fait délimité (virgule, point-virgule,
    tabulation, pipe) mais juste enregistré en .txt plutôt qu'en .csv —
    cas très fréquent (export brut d'un outil, copier-coller de tableur).
    """
    import csv, io

    text = "\n".join(lines)
    if not text.strip():
        return None
    try:
        dialect = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|")
    except csv.Error:
        return None

    rows = list(csv.reader(io.StringIO(text), dialect))
    rows = [r for r in rows if any(cell.strip() for cell in r)]
    if len(rows) < 2 or len(rows[0]) < 2:
        return None

    max_cols = max(len(r) for r in rows)
    rows = [r + [None] * (max_cols - len(r)) for r in rows]
    df = pd.DataFrame(rows)

    best_score = max(_score_header_row(df.iloc[i], column_mapping) for i in range(min(5, len(df))))
    if best_score <= 0:
        return None
    return df


def _try_fixed_width(lines: list[str], column_mapping: dict = None) -> pd.DataFrame | None:
    """
    Tentative n°2 : colonnes alignées par des espaces multiples, typique
    des rapports générés par des outils en ligne de commande ou des
    exports de systèmes legacy (ex. sortie brute d'un annuaire, rapport
    imprimé puis converti en texte).
    """
    rows = [re.split(r"\s{2,}", line.strip()) for line in lines if line.strip()]
    rows = [r for r in rows if len(r) >= 2]
    if len(rows) < 2:
        return None

    max_cols = max(len(r) for r in rows)
    rows = [r + [None] * (max_cols - len(r)) for r in rows]
    df = pd.DataFrame(rows)

    best_score = max(_score_header_row(df.iloc[i], column_mapping) for i in range(min(5, len(df))))
    if best_score <= 0:
        return None
    return df


def _try_key_value_blocks(raw_text: str, column_mapping: dict = None) -> pd.DataFrame | None:
    """
    Tentative n°3 : un enregistrement par bloc, séparé par des lignes
    vides, chaque ligne du bloc étant "clé: valeur" ou "clé= valeur"
    (avec ou sans puce). Format courant dans les comptes-rendus,
    fiches individuelles collées dans un document, ou exports type
    "un utilisateur par fiche".

    Exemple reconnu :
        Nom d'utilisateur: jdupont
        Système: Active Directory
        Statut: Actif

        Nom d'utilisateur: mfofana
        Système: CRM
        Statut: Actif
    """
    blocks = re.split(r"\n\s*\n", raw_text.strip())
    line_pattern = re.compile(r"^[-*•]?\s*([^:=]{2,60}?)\s*[:=]\s*(.+)$")

    records = []
    for block in blocks:
        record = {}
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            match = line_pattern.match(line)
            if match:
                key, value = match.group(1).strip(), match.group(2).strip()
                record[key] = value
        if len(record) >= 2:  # un bloc avec au moins 2 champs a une chance d'être un vrai enregistrement
            records.append(record)

    if len(records) < 1:
        return None

    df = pd.DataFrame(records)
    # Les clés extraites sont déjà les noms de colonnes réels (pas de ligne
    # d'en-tête à détecter séparément) : on vérifie juste qu'au moins une
    # d'entre elles est reconnaissable, pour éviter de valider n'importe
    # quel texte structuré par erreur.
    best_score = _score_header_row(pd.Series(df.columns), column_mapping)
    if best_score <= 0:
        return None
    return df


def _read_txt(path: Path, column_mapping: dict = None) -> tuple[pd.DataFrame, bool]:
    """
    Lit un fichier .txt en essayant plusieurs interprétations dans l'ordre
    de fiabilité décroissante, jusqu'à ce que l'une d'elles produise un
    résultat exploitable.

    Retourne (dataframe, header_already_named) : le second élément indique
    si les colonnes du DataFrame ont déjà leurs vrais noms (cas des blocs
    clé-valeur) ou si une détection d'en-tête classique reste à faire.
    """
    with open(path, encoding=_detect_encoding(path), errors="replace") as f:
        raw_text = f.read()

    lines = [l for l in raw_text.splitlines()]
    non_empty_lines = [l for l in lines if l.strip()]

    df = _try_delimited(non_empty_lines, column_mapping)
    if df is not None:
        logger.info("Fichier texte interprété comme des données délimitées.")
        return df, False

    df = _try_fixed_width(non_empty_lines, column_mapping)
    if df is not None:
        logger.info("Fichier texte interprété comme des colonnes alignées par espaces.")
        return df, False

    df = _try_key_value_blocks(raw_text, column_mapping)
    if df is not None:
        logger.info(f"Fichier texte interprété comme {len(df)} bloc(s) clé-valeur.")
        return df, True

    raise IngestionError(
        f"Impossible d'interpréter la structure de {path.name}. "
        "Formats texte reconnus : valeurs délimitées (virgule, point-virgule, "
        "tabulation, pipe), colonnes alignées par des espaces, ou blocs "
        "'clé: valeur' séparés par des lignes vides."
    )


def _read_docx(path: Path, column_mapping: dict = None) -> tuple[pd.DataFrame, bool]:
    """
    Lit un fichier Word. Essaie d'abord d'y trouver un tableau ; si aucun
    tableau n'est présent, retombe sur les mêmes stratégies de lecture de
    texte libre que pour un .txt, appliquées au texte des paragraphes.
    """
    try:
        from docx import Document
    except ImportError as e:
        raise IngestionError(
            "La bibliothèque 'python-docx' est requise pour lire les fichiers "
            ".docx. Installez-la avec : pip install python-docx"
        ) from e

    doc = Document(str(path))

    if doc.tables:
        # Un document peut contenir plusieurs tableaux légitimes (ex. un
        # par système/application, comme dans un rapport d'audit multi-
        # systèmes). Ne garder que "le meilleur" en perdrait silencieusement
        # tous les autres. On repère d'abord la ligne d'en-tête de référence
        # (meilleur score), puis on concatène tous les tableaux dont le
        # nombre de colonnes correspond — en excluant les répétitions
        # d'en-tête plutôt que de les traiter comme des données.
        best_header_row, best_score, ref_ncols = None, -1, None
        all_table_rows = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            all_table_rows.append(rows)
            table_df = pd.DataFrame(rows)
            for i in range(min(3, len(table_df))):
                score = _score_header_row(table_df.iloc[i], column_mapping)
                if score > best_score:
                    best_score = score
                    best_header_row = rows[i]
                    ref_ncols = len(rows[i])

        if best_header_row is not None:
            combined_rows = []
            for rows in all_table_rows:
                for row in rows:
                    if len(row) != ref_ncols:
                        continue
                    row_score = _score_header_row(pd.Series(row), column_mapping)
                    if row_score >= best_score * 0.9:
                        continue  # répétition de l'en-tête dans un autre tableau
                    combined_rows.append(row)

            logger.info(
                f"{len(doc.tables)} tableau(x) détecté(s) dans le document, "
                f"{len(combined_rows)} ligne(s) de données assemblées "
                f"(en-tête score={best_score})."
            )
            final_rows = [best_header_row] + combined_rows
            max_cols = max(len(r) for r in final_rows)
            final_rows = [r + [None] * (max_cols - len(r)) for r in final_rows]
            return pd.DataFrame(final_rows), False

    # Aucun tableau exploitable : on retombe sur le texte des paragraphes
    logger.info("Aucun tableau exploitable — tentative de lecture en texte libre.")
    paragraphs_with_blanks = [p.text for p in doc.paragraphs]
    non_empty = [p for p in paragraphs_with_blanks if p.strip()]

    df = _try_delimited(non_empty, column_mapping)
    if df is not None:
        logger.info("Contenu du document interprété comme des données délimitées.")
        return df, False

    df = _try_fixed_width(non_empty, column_mapping)
    if df is not None:
        logger.info("Contenu du document interprété comme des colonnes alignées.")
        return df, False

    df = _try_key_value_blocks("\n".join(paragraphs_with_blanks), column_mapping)
    if df is not None:
        logger.info(f"Contenu du document interprété comme {len(df)} bloc(s) clé-valeur.")
        return df, True

    raise IngestionError(
        f"Aucun tableau ni structure de données reconnaissable dans {path.name}. "
        "Formats reconnus : tableau Word, texte délimité, colonnes alignées, "
        "ou blocs 'clé: valeur' séparés par des lignes vides."
    )



def _read_json(path: Path) -> pd.DataFrame:
    """
    Lit un fichier JSON. Accepte :
        - une liste d'objets : [{"username": "...", ...}, ...]
        - un objet unique contenant une liste sous une clé courante
          (results, data, users, accounts, records, items, value)
        - un objet unique représentant un seul compte
    """
    import json

    with open(path, encoding=_detect_encoding(path)) as f:
        data = json.load(f)

    if isinstance(data, list):
        records = data
    elif isinstance(data, dict):
        list_keys = ["results", "data", "users", "accounts", "records", "items", "value"]
        records = None
        for key in list_keys:
            if key in data and isinstance(data[key], list):
                records = data[key]
                break
        if records is None:
            # objet unique = un seul enregistrement
            records = [data]
    else:
        raise IngestionError(f"Structure JSON non reconnue dans {path.name}.")

    if not records:
        raise IngestionError(f"Aucun enregistrement trouvé dans {path.name}.")

    # Aplatit les dictionnaires imbriqués simples (ex. {"user": {"name": "..."}})
    df = pd.json_normalize(records, sep="_")
    return df


def _read_xml(path: Path) -> pd.DataFrame:
    """
    Lit un fichier XML. Suppose une structure répétitive classique
    (une balise par enregistrement, ex. <user>...</user> ou <account>...</account>),
    avec les champs en sous-balises ou en attributs.
    """
    try:
        df = pd.read_xml(path)
    except Exception as e:
        raise IngestionError(
            f"Impossible d'interpréter la structure XML de {path.name} ({e}). "
            "Le fichier doit contenir des éléments répétitifs représentant "
            "chacun un compte (ex. <user>...</user>)."
        ) from e

    if df.empty:
        raise IngestionError(f"Aucun enregistrement exploitable trouvé dans {path.name}.")
    return df


def _read_html(path: Path, column_mapping: dict = None) -> pd.DataFrame:
    """
    Lit un fichier HTML contenant un ou plusieurs tableaux (ex. export copié
    depuis une page web/intranet). Garde le tableau le plus pertinent, même
    logique que pour les tableaux Word multiples.

    pandas détecte déjà l'en-tête via les balises <th>/<thead> : on
    reconstruit une forme "brute" (en-tête recollé comme première ligne)
    pour rester cohérent avec les autres lecteurs et laisser la détection
    d'en-tête standard s'appliquer une seule fois, au bon endroit.
    """
    try:
        tables = pd.read_html(path)
    except ValueError as e:
        raise IngestionError(f"Aucun tableau trouvé dans {path.name} ({e}).") from e

    best_raw_rows, best_score = None, -1
    for table_df in tables:
        raw_rows = [list(table_df.columns)] + table_df.astype(object).values.tolist()
        candidate = pd.DataFrame(raw_rows)
        score = max(_score_header_row(candidate.iloc[i], column_mapping) for i in range(min(3, len(candidate))))
        if score > best_score:
            best_raw_rows, best_score = raw_rows, score

    if best_raw_rows is None:
        raise IngestionError(f"Aucun tableau exploitable trouvé dans {path.name}.")

    logger.info(f"{len(tables)} tableau(x) HTML détecté(s), le plus pertinent retenu (score={best_score}).")
    return pd.DataFrame(best_raw_rows)


_UAC_ACCOUNTDISABLE_BIT = 0x2  # bit standard Active Directory pour "compte désactivé"


def _read_ldif(path: Path) -> pd.DataFrame:
    """
    Lit un fichier LDIF (export natif LDAP/Active Directory).

    Structure LDIF : des entrées séparées par des lignes vides, chaque
    ligne étant "attribut: valeur" (ou "attribut:: valeur_base64" pour les
    valeurs encodées). Un attribut peut apparaître plusieurs fois (valeurs
    multiples) — on garde alors la première occurrence pour rester simple.

    Cas particulier traité : 'userAccountControl' est un bitmask numérique
    (pas un statut texte). On le décode ici pour en tirer directement un
    statut Active/Disabled exploitable par le reste du pipeline.
    """
    with open(path, encoding=_detect_encoding(path), errors="replace") as f:
        raw_text = f.read()

    # Les lignes de continuation LDIF commencent par un espace : elles
    # prolongent la ligne précédente et doivent être recollées avant parsing.
    unfolded_lines = []
    for line in raw_text.splitlines():
        if line.startswith(" ") and unfolded_lines:
            unfolded_lines[-1] += line[1:]
        else:
            unfolded_lines.append(line)

    entries_text = re.split(r"\n\s*\n", "\n".join(unfolded_lines))
    records = []

    for entry_text in entries_text:
        record = {}
        for line in entry_text.split("\n"):
            line = line.rstrip()
            if not line or line.startswith("#"):
                continue
            # "attribut:: valeur" = base64, on ignore le décodage (rare
            # pour les champs texte qui nous intéressent ici)
            match = re.match(r"^([\w;-]+)::?\s*(.*)$", line)
            if not match:
                continue
            key, value = match.group(1), match.group(2).strip()
            if key.lower() == "useraccountcontrol":
                try:
                    flags = int(value)
                    value = "Disabled" if (flags & _UAC_ACCOUNTDISABLE_BIT) else "Active"
                except ValueError:
                    pass
            if key not in record:  # garde la première valeur si attribut répété
                record[key] = value
        if record:
            records.append(record)

    if not records:
        raise IngestionError(
            f"Aucune entrée LDIF exploitable trouvée dans {path.name}. "
            "Format attendu : entrées séparées par des lignes vides, "
            "lignes 'attribut: valeur'."
        )

    # Un export LDIF représente par nature un seul système (l'annuaire
    # LDAP/AD lui-même) : il n'y a jamais de colonne "système" explicite
    # dans les données, contrairement à un export multi-systèmes. On
    # l'ajoute donc nous-mêmes plutôt que d'échouer sur un champ obligatoire
    # qui n'a structurellement aucune raison d'exister dans ce format.
    for record in records:
        record.setdefault("system", "Active Directory / LDAP")

    logger.info(f"{len(records)} entrée(s) LDIF décodée(s).")
    return pd.DataFrame(records)


def _read_pdf(path: Path, column_mapping: dict = None) -> pd.DataFrame:
    """
    Extrait un tableau depuis un PDF, sur l'ensemble de ses pages.

    Un tableau réel s'étale très souvent sur plusieurs pages (un export de
    800 lignes ne tient jamais sur une seule page) : l'en-tête n'apparaît
    en général qu'une fois, en haut de la première page, parfois répété en
    haut de chaque page suivante. On identifie d'abord la page contenant le
    véritable en-tête (meilleur score de reconnaissance de colonnes), puis
    on concatène les lignes de TOUTES les pages dont le nombre de colonnes
    correspond à ce même tableau — en ignorant les répétitions de l'en-tête
    sur les pages suivantes plutôt que de les traiter comme des données.
    """
    try:
        import pdfplumber
    except ImportError as e:
        raise IngestionError(
            "La bibliothèque 'pdfplumber' est requise pour lire les PDF. "
            "Installez-la avec : pip install pdfplumber"
        ) from e

    all_tables = []  # toutes les tables trouvées, toutes pages confondues, dans l'ordre
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            if not tables:
                logger.warning(
                    "Aucune bordure de tableau détectée sur une page : tentative par "
                    "alignement du texte, moins fiable (peut mal découper "
                    "des colonnes ou des lignes) — vérifiez le résultat."
                )
                tables = page.extract_tables(
                    table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                    }
                )
            all_tables.extend(t for t in tables if t)

    if not all_tables:
        raise IngestionError(
            f"Aucun tableau détecté dans {path.name}. L'extraction de tableaux "
            "PDF est fiable uniquement si le PDF contient une vraie grille "
            "(pas une image scannée ni une mise en page libre)."
        )

    # 1) Repérer la ligne d'en-tête de référence : le meilleur score, tous
    #    tableaux/pages confondus.
    best_header_row, best_score, ref_ncols = None, -1, None
    for table in all_tables:
        table_df = pd.DataFrame(table)
        for i in range(min(3, len(table_df))):
            score = _score_header_row(table_df.iloc[i], column_mapping)
            if score > best_score:
                best_score = score
                best_header_row = table[i]
                ref_ncols = len(table[i])

    if best_header_row is None:
        raise IngestionError(
            f"Aucun en-tête reconnaissable dans les tableaux de {path.name}."
        )

    # 2) Concaténer les lignes de TOUTES les pages ayant le même nombre de
    #    colonnes que la référence (les petits tableaux annexes, type
    #    bloc de signature, ont généralement un nombre de colonnes
    #    différent et sont donc naturellement exclus). On saute les lignes
    #    qui ressemblent fortement à une répétition de l'en-tête (score
    #    proche du meilleur score) plutôt que de les garder comme données.
    combined_rows = []
    for table in all_tables:
        for row in table:
            if len(row) != ref_ncols:
                continue
            row_score = _score_header_row(pd.Series(row), column_mapping)
            if row_score >= best_score * 0.9:
                continue  # répétition de l'en-tête sur cette page, pas une donnée
            combined_rows.append(row)

    logger.info(
        f"Tableau PDF assemblé sur {len(all_tables)} fragment(s) de page(s) : "
        f"{len(combined_rows)} ligne(s) de données (en-tête score={best_score})."
    )

    rows = [best_header_row] + combined_rows
    max_cols = max(len(r) for r in rows)
    rows = [list(r) + [None] * (max_cols - len(r)) for r in rows]
    return pd.DataFrame(rows)


# Extensions traitées nativement par _load_single_file (utilisé aussi par
# _read_zip pour savoir quels fichiers internes tenter d'ouvrir).
SUPPORTED_EXTENSIONS = [
    ".csv", ".xlsx", ".xls", ".docx", ".txt", ".json", ".xml", ".html", ".htm",
    ".ldif", ".pdf",
]


def _load_single_file(
    path: Path, column_mapping: dict = None, required_fields: list = None,
    default_system: str | None = None,
) -> pd.DataFrame:
    """Charge un unique fichier (tous formats sauf .zip) et retourne un DataFrame standardisé."""
    logger.info(f"Lecture du fichier : {path.name}")

    header_already_named = False
    suffix = path.suffix.lower()

    if suffix in [".xlsx", ".xls"]:
        raw = pd.read_excel(path, header=None, sheet_name=0)
    elif suffix == ".csv":
        raw = _read_ragged_csv(path)
    elif suffix == ".docx":
        raw, header_already_named = _read_docx(path, column_mapping)
    elif suffix == ".txt":
        raw, header_already_named = _read_txt(path, column_mapping)
    elif suffix == ".json":
        raw = _read_json(path)
        header_already_named = True
    elif suffix == ".xml":
        raw = _read_xml(path)
        header_already_named = True
    elif suffix in [".html", ".htm"]:
        raw = _read_html(path, column_mapping)
    elif suffix == ".ldif":
        raw = _read_ldif(path)
        header_already_named = True
    elif suffix == ".pdf":
        raw = _read_pdf(path, column_mapping)
    else:
        raise IngestionError(
            f"Format de fichier non supporté : {path.suffix}. "
            f"Formats acceptés : {', '.join(SUPPORTED_EXTENSIONS)}, .zip"
        )

    if header_already_named:
        df = raw.reset_index(drop=True)
    else:
        header_row_idx = _detect_header_row(raw, column_mapping)
        df = raw.iloc[header_row_idx + 1:].copy()
        df.columns = raw.iloc[header_row_idx]
        df = df.dropna(how="all").reset_index(drop=True)

    df = standardize_columns(df, column_mapping)

    # Un export "brut" d'un seul système (ex. extraction Active Directory
    # pure) ne contient souvent aucune colonne identifiant le système lui-
    # même : cette information est implicite (tout le fichier = ce système),
    # pas une donnée par ligne. Plutôt que d'échouer, on comble ce vide :
    # priorité au nom explicite passé à l'appel (ex. depuis le dashboard),
    # sinon repli automatique sur le nom du fichier — jamais d'échec pour
    # cette seule raison. Uniquement si 'system' est effectivement requis
    # pour ce domaine (inutile, par ex., pour un référentiel RH).
    effective_required = required_fields if required_fields is not None else REQUIRED_FIELDS
    if "system" not in df.columns and "system" in effective_required:
        resolved_system = default_system or path.stem
        df["system"] = resolved_system
        logger.info(f"Colonne 'system' absente du fichier : valeur par défaut appliquée ('{resolved_system}').")

    validate_required_fields(df, required_fields)

    logger.info(f"Ingestion réussie : {len(df)} lignes, colonnes finales : {list(df.columns)}")
    return df


def _read_zip(
    path: Path, column_mapping: dict = None, required_fields: list = None,
    default_system: str | None = None,
) -> pd.DataFrame:
    """
    Extrait une archive ZIP et traite chaque fichier supporté qu'elle
    contient, puis concatène tous les résultats. Utile pour un export
    mensuel regroupant plusieurs systèmes (un fichier par système) dans
    une seule archive.

    Les fichiers dans un format non supporté ou illisibles sont ignorés
    avec un avertissement, plutôt que de faire échouer tout le traitement.

    Si un fichier interne n'a pas de colonne 'system', son propre nom de
    fichier sert de valeur par défaut (plus pertinent que `default_system`
    partagé, vu qu'un ZIP regroupe typiquement un système par fichier) —
    sauf si `default_system` est explicitement fourni, auquel cas il
    s'applique à tous les fichiers de l'archive.
    """
    import zipfile
    import tempfile

    dfs = []
    skipped = []

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_dir_path = Path(tmp_dir)
        with zipfile.ZipFile(path) as zf:
            zf.extractall(tmp_dir_path)

        candidate_files = sorted(
            p for p in tmp_dir_path.rglob("*")
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
            and not p.name.startswith(".") and "__MACOSX" not in str(p)
        )

        if not candidate_files:
            raise IngestionError(
                f"Aucun fichier de format supporté trouvé dans l'archive {path.name}."
            )

        for f in candidate_files:
            try:
                df = _load_single_file(f, column_mapping, required_fields, default_system=default_system)
                df["_source_file"] = f.name
                dfs.append(df)
            except IngestionError as e:
                skipped.append((f.name, str(e)))
                logger.warning(f"Fichier ignoré dans l'archive ({f.name}) : {e}")

    if not dfs:
        raise IngestionError(
            f"Aucun fichier exploitable dans l'archive {path.name}. "
            f"Fichiers trouvés mais ignorés : {[s[0] for s in skipped]}"
        )

    logger.info(f"{len(dfs)} fichier(s) traité(s) avec succès dans l'archive (sur {len(candidate_files)}).")
    combined = pd.concat(dfs, ignore_index=True, sort=False)
    return combined


def load_file(path: str | Path, default_system: str | None = None) -> pd.DataFrame:
    """
    Point d'entrée standard : charge un fichier d'export IAM/accès, avec
    le référentiel de colonnes par défaut (config/column_mapping.py).

    `default_system` : nom de système à appliquer si le fichier ne contient
    aucune colonne l'identifiant lui-même (cas fréquent d'un export brut
    d'un seul système). Sans valeur fournie, le nom du fichier sert de
    repli automatique — la fonction ne lève jamais d'erreur pour ce seul
    motif.
    """
    path = Path(path)
    if not path.exists():
        raise IngestionError(f"Fichier introuvable : {path}")

    if path.suffix.lower() == ".zip":
        logger.info(f"Lecture de l'archive : {path.name}")
        return _read_zip(path, default_system=default_system)

    return _load_single_file(path, default_system=default_system)


def load_file_with_mapping(
    path: str | Path, column_mapping: dict, required_fields: list,
    default_system: str | None = None,
) -> pd.DataFrame:
    """
    Variante de load_file() pour un domaine différent de celui des exports
    d'accès (ex. un export RH), avec son propre référentiel de colonnes et
    ses propres champs obligatoires. Réutilise exactement la même logique
    de lecture universelle (tous formats, détection d'en-tête, etc.).
    """
    path = Path(path)
    if not path.exists():
        raise IngestionError(f"Fichier introuvable : {path}")

    if path.suffix.lower() == ".zip":
        logger.info(f"Lecture de l'archive : {path.name}")
        return _read_zip(path, column_mapping, required_fields, default_system=default_system)

    return _load_single_file(path, column_mapping, required_fields, default_system=default_system)


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage : python ingest.py <chemin_fichier>")
        sys.exit(1)
    print(load_file(sys.argv[1]).head())
